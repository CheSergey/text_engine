import docx
import docx.shared
import openai as opa
import dotenv as dtv

dtv.load_dotenv(".env", override=True)
api_key = os.getenv("OPENAI_API_KEY")
client = opa.OpenAI(api_key = api_key)

def main(settings, client):
    main_choice = show_main_menu()
    if main_choice == 1:
        translate_document_flow(settings, client)
    elif main_choice == 2:
        correction_document_flow()
    elif main_choice == 3:
        format_document_flow(settings)
    elif main_choice == 4:
        GPT_settings_flow()
    elif main_choice == 5:
        print("Выход")
        return False
    return True

def show_main_menu():
    while True:
        choice = input("\n=== Lingua Engine ===\n\n1. Перевести документ\n2. Проверить орфографию\n3. Форматирование документа\n4. Настройки модели GPT\n5. Выход\n\nВаш выбор:").strip() # Вызов основного меню
        try:
            choice = int(choice)
        except:
            print("\n*** Это не число. Выберите число. ***")
            continue
        if not (1 <= choice <= 5):
            print("*** Введите число от 1 до 5 ***")
            continue
        return choice

def translate_document_flow(settings, client):
    while True:
        translation_menu_choice = show_translation_menu()
        if translation_menu_choice == 1:
            text = try_to_open()
            if not text:
                return True
            prompt_text = try_to_open_prompt()
            if not prompt_text:
                return True
            blocks = process_text(text)
            chunks = chunk_creation(blocks, settings["max_chars"])
            if settings["debug_mode"]:
                chunk_saver(chunks)
            translated_chunks = translator(chunks, settings, client, prompt_text)
            text_for_export = export_text_preparation_translate(translated_chunks)
            output_file_name = get_output_file_name()
            if output_file_name is None:
                return None
            if output_file_name.endswith(".txt"):
                translated_document_saver(text_for_export, output_file_name)
            if output_file_name.endswith(".docx"):
                document = export_document_format(text_for_export, settings)
                translated_document_saver_docx(document, output_file_name)
            return None
        elif translation_menu_choice == 2:
            translation_menu_settings_workflow(settings)
        elif translation_menu_choice == 3:
            return
            
def show_translation_menu():
    while True:
        choice = input("\n=== Lingua Engine ===\n\n1. Загрузить документ\n2. Дополнительные настройки\n3. Назад\n\nВаш выбор:").strip() # чанк - это кусок текста 
        try:
            choice = int(choice)
        except:
            print("\n*** Это не число. Выберите число. ***")
            continue
        if not (1 <= choice <= 3):
            print("*** Введите число от 1 до 3 ***")
            continue
        return choice
    
def translation_menu_settings_workflow(settings):
    while True:
        choice = show_translation_menu_settings(settings)
        if choice == 11:
            return
        apply_translate_settings(settings, choice)

def show_translation_menu_settings(settings):
    while True:
        choice = input(

    f"\n=== Lingua Engine ===\n\n"

    f"1. Язык:{settings['language']}\n"
    f"2. Дебаггинг:{settings['debug_mode']}\n"
    f"3. Объем чанка:{settings['max_chars']}\n"
    f"4. Модель:{settings['model']}\n"
    f"5. Temperature:{settings['temperature']}\n"
    f"6. Межстрочный интервал:{settings['line_spacing']}\n"
    f"7. Интервал до:{settings['space_before']}\n"
    f"8. Интервал после:{settings['space_after']}\n"
    f"9. Шрифт:{settings['font_name']}\n"
    f"10. Размер шрифта:{settings['font_size']}\n"
    f"11. Назад\n\n"

    f"Ваш выбор: "
).strip()
        try:
            choice = int(choice)
        except:
            print("\n*** Это не число. Выберите число. ***")
            continue
        if not (1 <= choice <= 11):
            print("*** Введите число от 1 до 11 ***")
            continue
        return choice


settings = {
    "debug_mode": False,
    "max_chars": 3000,
    "language": "English",
    "model": "gpt-5.5",
    "temperature": 0.2,
    "space_before": 0,
    "space_after": 0,
    "font_name": "Times New Roman",
    "font_size": 12,
    "line_spacing": 1.15
}

def apply_translate_settings(settings, choice):
    if choice == 1:
        language_choice = show_language_options(settings)
        if language_choice == 1: 
            settings["language"] = "English"
        if language_choice == 2:
            settings["language"] = "Czech"
        if language_choice == 3:
            return settings
    if choice == 2:
        if settings["debug_mode"] == True:
            settings["debug_mode"] = False
        else:
            settings["debug_mode"] = True
    if choice == 3:
        while True:    
            max_chars = input("Введите количество символов для одного чанка (от 1000 до 20000 символов):").strip() # чанк - это кусок текста 
            try:
                max_chars = int(max_chars)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (1000 <= max_chars <= 20000):
                print("*** Введите число от 1000 до 20000 ***")
                continue
            settings["max_chars"] = max_chars
            break
    if choice == 4:
        while True:
            model = input("Введите название модели (например: gpt-5.5): ").strip()
            if model == "":
                print("*** Вы ничего не ввели ***")
                continue
            if not model.startswith("gpt-"):
                print("*** Модель должна начинаться с 'gpt-' ***")
                continue
            settings["model"] = model
            break
    if choice == 5: 
        while True:    
            temperature = input("Введите параметр temperature (от 0 до 2):").strip()
            try:
                temperature = float(temperature)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= temperature <= 2):
                print("*** Введите число от 0 до 2 ***")
                continue
            settings["temperature"] = temperature
            break
    if choice == 6:
        while True:    
            line_spacing = input("Введите межстрочный интервал (от 0 до 3):").strip()
            try:
                line_spacing = float(line_spacing)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= line_spacing <= 3):
                print("*** Введите число от 0 до 3 ***")
                continue
            settings["line_spacing"] = line_spacing
            break
    if choice == 7:
        while True:    
            space_before = input("Введите отступы до (от 0 до 10):").strip()
            try:
                space_before = int(space_before)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= space_before <= 10):
                print("*** Введите число от 0 до 10 ***")
                continue
            settings["space_before"] = space_before
            break
    if choice == 8:
        while True:    
            space_after = input("Введите отступы после (от 0 до 10):").strip()
            try:
                space_after = int(space_after)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= space_after <= 10):
                print("*** Введите число от 0 до 10 ***")
                continue
            settings["space_after"] = space_after
            break
    if choice == 9:
        while True:    
            font_name = input("Введите название шрифта:").strip()
            try:
                font_name = str(font_name)
            except:
                print("*** Это не строка. Введите название. ***")
                continue
            settings["font_name"] = font_name
            break
    if choice == 10:
        while True:
            font_size = input("Введите размер шрифта:").strip()  
            try:
                font_size = int(font_size)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            settings["font_size"] = font_size
            break
    return settings

def show_language_options(settings):
    while True:
        choice = input(
        f"\n=== Lingua Engine ===\n\n"

        f"Текущий язык:{settings["language"]}\n\n"

        "Выберите язык:\n1. English\n2. Czech\n3. Назад\n\n"

        "Ваш выбор:"
).strip()
        try:
            choice = int(choice)
        except:
            print("\n*** Это не число. Выберите число. ***")
            continue
        if choice == 3:
            return choice
        if not (1 <= choice <= 3):
            print("*** Введите число от 1 до 3 ***")
            continue
        return choice

def correction_document_flow():
    pass

def correction_menu_flow():
    pass

def show_correction_menu():
    pass

def format_document_flow(settings):
    while True:
        format_menu_choice = format_document_menu()
        if format_menu_choice == 1:
            text = try_to_open()
            if not text:
                return True
            blocks = process_text(text)
            text_for_export = export_text_preparation_format(blocks)
            output_file_name = get_output_file_name()
            if output_file_name is None:
                return None
            if output_file_name.endswith(".txt"):
                translated_document_saver(text_for_export, output_file_name)
            if output_file_name.endswith(".docx"):
                document = export_document_format(text_for_export, settings)
                translated_document_saver_docx(document, output_file_name)
            return None
        elif format_menu_choice == 2:
            format_menu_settings_workflow(settings)
        elif format_menu_choice == 3:
            return
        
def format_menu_settings_workflow(settings):
    while True:
        choice = show_format_options(settings)
        if choice == 6:
            return
        apply_format_settings(settings, choice)

def format_document_menu():
    while True:
        choice = input(
        f"\n=== Lingua Engine ===\n\n"

        f"1. Загрузить документ\n"
        f"2. Дополнительные настройки.\n"
        f"3. Назад\n\n"

        "Ваш выбор:"
).strip()
        try:
            choice = int(choice)
        except:
            print("\n*** Это не число. Выберите число. ***")
            continue
        if not (1 <= choice <= 3):
            print("*** Введите число от 1 до 3 ***")
            continue
        return choice

def show_format_options(settings):
    while True:
        choice = input(
    f"\n=== Lingua Engine ===\n\n"

    f"1. Межстрочный интервал: {settings['line_spacing']}\n"
    f"2. Интервал до: {settings['space_before']}\n"
    f"3. Интервал после: {settings['space_after']}\n"
    f"4. Шрифт: {settings['font_name']}\n"
    f"5. Размер шрифта: {settings['font_size']}\n"
    "6. Назад\n\n"

    "Ваш выбор:"
).strip()
        try:
            choice = int(choice)
        except:
            print("\n*** Это не число. Выберите число. ***")
            continue
        if not (1 <= choice <= 6):
            print("*** Введите число от 1 до 6 ***")
            continue
        return choice

def apply_format_settings(settings, choice):
    if choice == 1:
        while True:    
            line_spacing = input("Введите межстрочный интервал (от 0 до 3):").strip()
            try:
                line_spacing = float(line_spacing)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= line_spacing <= 3):
                print("*** Введите число от 0 до 3 ***")
                continue
            settings["line_spacing"] = line_spacing
            break
    if choice == 2:
        while True:    
            space_before = input("Введите отступы до (от 0 до 10):").strip()
            try:
                space_before = int(space_before)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= space_before <= 10):
                print("*** Введите число от 0 до 10 ***")
                continue
            settings["space_before"] = space_before
            break
    if choice == 3:
        while True:    
            space_after = input("Введите отступы после (от 0 до 10):").strip()
            try:
                space_after = int(space_after)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            if not (0 <= space_after <= 10):
                print("*** Введите число от 0 до 10 ***")
                continue
            settings["space_after"] = space_after
            break
    if choice == 4:
        while True:    
            font_name = input("Введите название шрифта:").strip()
            try:
                font_name = str(font_name)
            except:
                print("*** Это не строка. Введите название. ***")
                continue
            settings["font_name"] = font_name
            break
    if choice == 5:
        while True:    
            font_size = input("Введите размер шрифта:").strip()
            try:
                font_size = int(font_size)
            except:
                print("*** Это не число. Введите число. ***")
                continue
            settings["font_size"] = font_size
            break
    return settings

def GPT_settings_flow():
    pass

def GPT_settings_menu():
    pass

def getting_file_name_input():
    while True:
        file_name_input = input("Введите название файла с текстом для перевода, включительно формата (.txt) или (.docx). Для выхода введите (esc): ").strip()
        if file_name_input.lower() == "esc":
            return None
        elif not file_name_input.lower().endswith(".docx") and not file_name_input.lower().endswith(".txt"):
            print("Вы не ввели название или формат")
            continue
        else:
            return file_name_input

def getting_translation_prompt_name():
    while True:
        prompt_name_input = input("Введите название файла с промтом, включительно формата (.txt). Для выхода введите (esc): ").strip()
        if prompt_name_input.lower() == "esc":
            return None
        elif not prompt_name_input.lower().endswith(".txt"):
            print("Введите название с форматом .txt")
            continue
        else:
            return prompt_name_input

def open_file(file_name):
    if file_name == "":
        text = ""
        error = "Имя файла не задано"
        print(error)
        return text, error       
    try:
        with open(file_name,"r",encoding="utf-8") as file:
            text = file.read()
            error = ""
            return text, error
    except FileNotFoundError:
        error = "Файл не найден"
        text = ""
        print (error)
    return text, error 

def open_file_docx(file_name):
    if file_name == "":
        document = ""
        error = "Имя файла не задано"
        print(error)
        return document, error
    try:
        document = docx.Document(file_name)
        error = ""
    except docx.opc.exceptions.PackageNotFoundError: # except Exception as e: - более униврсельный способ поиска ошибок
        error = "Файл не найден" 
        document = ""
        print(error)  # альтернатива, чтобы узнать тип ошибки print(type(e))
    return document, error

def try_to_open():
    file_name = getting_file_name_input()
    if not file_name:
        text = ""
    elif file_name.lower().endswith(".txt"):
        text, error = open_file(file_name)
        if error:
            text = ""
    elif file_name.lower().endswith(".docx"):
        document, error = open_file_docx(file_name)
        if error:
            text = ""
        else:
            text = extract_text_from_docx(document)
    return text

def try_to_open_prompt():
    prompt_name = getting_translation_prompt_name()
    if not prompt_name:
        prompt_text = ""
    elif prompt_name.lower().endswith(".txt"):
        prompt_text, error= open_file(prompt_name)
        if error:
            prompt_text = ""
    return prompt_text

def extract_text_from_docx(document):
    text = ""
    for paragraph in document.paragraphs:
        if text == "":
            text = paragraph.text
        else:
            text += f"\n{paragraph.text}"
    return text

def process_text(text):
    paragraphs = text_divider(text)    
    blocks = addition_divider(paragraphs)
    blocks = semantic_block_builder(blocks)
    return blocks

def text_divider(text):
    paragraphs = text.split("\n\n")
    return paragraphs

def addition_divider(paragraphs):
    blocks = []
    for paragraph in paragraphs:
        if "\n" in paragraph:
            lines = paragraph.split("\n")
            for line in lines:
                if line.strip():
                    blocks.append(line)
        else: 
            if paragraph.strip():
                blocks.append(paragraph)
    return blocks

def semantic_block_builder(blocks):
    new_blocks = []
    temporary_text = ""
    for block in blocks:
        first_char = block.strip()[0]
        if first_char.islower():
            temporary_text += f"{block}\n" 
        else:    
            if temporary_text:
                temporary_text = f"\n{temporary_text}"
                new_blocks.append(temporary_text)
                temporary_text = ""
            new_blocks.append(block)
    if temporary_text:
        temporary_text = f"\n{temporary_text}"
        new_blocks.append(temporary_text)
    return new_blocks

def chunk_creation(blocks, max_chars):
    chunk = ""
    chunks = []
    for block in blocks:
        if len(chunk + block) <= max_chars and chunk == "":
            chunk = block
        elif len(chunk + "\n" + block) <= max_chars and chunk != "":
            chunk += "\n" + block
        else:
            chunks.append(chunk)
            chunk = block 
    if not chunk == "":
        chunks.append(chunk)     
    return chunks

def translator(chunks, settings, client, prompt_text):
    translated_chunks = []
    for i in range(len(chunks)):
        chunk = chunks[i]
        print(f"Перевожу чанк {i+1} из {len(chunks)}")
        response = client.responses.create(
            model = settings["model"],
            input = f"""TARGET LANGUAGE:
            {settings['language']}

            INSTRUCTIONS: 
            {prompt_text}
                    
            TEXT TO TRANSLATE:
            {chunk}
            """
            )
        translated_chunk = response.output_text
        translated_chunks.append(translated_chunk)
    return translated_chunks

def export_text_preparation_translate(translated_chunks):
    text = "\n".join(translated_chunks)
    return text

def export_text_preparation_format(blocks):
    text = "\n".join(blocks)
    return text

def export_document_format(text, settings):    
    document = docx.Document()
    paragraphs = text.split("\n")
    paragraphs_modified = []
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i].strip()
        if i > 0:
            previous = paragraphs[i-1].strip()
        else:
            previous = ""
        if i < len(paragraphs) - 1:
            following = paragraphs[i+1].strip()
        else:
            following = ""
        if paragraph.startswith("—") and previous and previous[0].isupper() and following and following[0] != "—":
            paragraph = f"\n{paragraph}\n"
            paragraphs_modified.append(paragraph)
        elif paragraph.startswith("—") and previous and previous[0].isupper():
            paragraph = f"\n{paragraph}"
            paragraphs_modified.append(paragraph)
        elif paragraph.startswith("—") and following and following[0] != "—":
            paragraph = f"{paragraph}\n"
            paragraphs_modified.append(paragraph)
        else:
            paragraphs_modified.append(paragraph)
    for paragraph in paragraphs_modified:
        doc_paragraph = document.add_paragraph(paragraph)
        doc_paragraph.paragraph_format.space_before = docx.shared.Pt(settings["space_before"])
        doc_paragraph.paragraph_format.space_after = docx.shared.Pt(settings["space_after"])
        doc_paragraph.paragraph_format.line_spacing = settings["line_spacing"]
        if doc_paragraph.runs:
            text_part = doc_paragraph.runs[0]
            text_part.font.name = settings["font_name"]
            text_part.font.size = docx.shared.Pt(settings["font_size"])
    return document

def get_output_file_name():
    while True:
        output_file_name = input("Введите название файла для сохранения результата, включительно формата (.txt / .docx) или ESC для выхода: ").strip()
        if not output_file_name:
            print("Вы не ввели название файла. ")
            continue
        if output_file_name.lower() == "esc":
            return None
        return output_file_name  

def translated_document_saver(text, output_file_name):
    with open(output_file_name, "w", encoding="utf-8") as result_file:
        result_file.write(text)
    return None

def translated_document_saver_docx(document, output_file_name):
    document.save(output_file_name)

def chunk_saver(chunks):
    number_of_chunk = 1
    for chunk in chunks:
        name_of_chunk_file = (f"chunk_{number_of_chunk}.txt")
        with open(name_of_chunk_file, "w", encoding="utf-8") as result_file:
            result_file.write(chunk)
            number_of_chunk += 1
    return None

while True:
    should_continue = main(settings, client)
    if not should_continue:
        break